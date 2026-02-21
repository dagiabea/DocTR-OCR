# Copyright (C) 2021-2025, Mindee.
# This program is licensed under the Apache License 2.0.
# See LICENSE or go to <https://opensource.org/licenses/Apache-2.0> for full license details.

import io
from typing import Any

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract plain text directly from DOCX file, bypassing OCR.
    
    Args:
        file_content: Binary content of DOCX file
        
    Returns:
        Extracted text as string
    """
    if not DOCX_AVAILABLE or DocxDocument is None:
        import sys
        python_path = sys.executable
        raise ValueError(
            f"python-docx is required for DOCX support. "
            f"Install it with: {python_path} -m pip install python-docx\n"
            f"Current Python: {python_path}"
        )
    
    try:
        doc = DocxDocument(io.BytesIO(file_content))
    except Exception as e:
        raise ValueError(f"Failed to open DOCX file: {str(e)}")
    
    # Extract all text from paragraphs (including empty ones for structure)
    full_text = []
    para_count = 0
    for paragraph in doc.paragraphs:
        para_count += 1
        para_text = paragraph.text.strip()
        # Include paragraph even if empty to preserve structure
        full_text.append(para_text if para_text else '')
    
    # Extract text from tables
    table_count = 0
    for table in doc.tables:
        table_count += 1
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Extract text from all paragraphs in cell
                cell_paras = []
                for para in cell.paragraphs:
                    cell_text = para.text.strip()
                    if cell_text:
                        cell_paras.append(cell_text)
                cell_text = ' '.join(cell_paras) if cell_paras else ''
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_rows.append(' | '.join(row_text))
        if table_rows:
            full_text.append('')  # Add blank line before table
            full_text.extend(table_rows)
            full_text.append('')  # Add blank line after table
    
    # Also try to extract from headers and footers
    try:
        for section in doc.sections:
            if section.header:
                header_text = []
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text.strip())
                if header_text:
                    full_text.insert(0, '\n'.join(header_text))
                    full_text.insert(1, '')  # Blank line after header
            
            if section.footer:
                footer_text = []
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text.strip())
                if footer_text:
                    full_text.append('')  # Blank line before footer
                    full_text.extend(footer_text)
    except Exception:
        # Headers/footers extraction is optional, continue if it fails
        pass
    
    # Try to extract from text boxes and shapes (if present)
    try:
        # Access document body XML to find text in text boxes
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.ns import qn
        
        body = doc.element.body
        for element in body.iter():
            if element.tag.endswith('}p'):  # Paragraph element
                para_text = ''
                for t in element.iter():
                    if t.tag.endswith('}t'):  # Text element
                        if t.text:
                            para_text += t.text
                if para_text.strip():
                    full_text.append(para_text.strip())
    except Exception:
        # Text box extraction is optional
        pass
    
    # Join all text
    text_content = '\n'.join(full_text)
    
    # Remove excessive blank lines but keep structure
    lines = text_content.split('\n')
    cleaned_lines = []
    prev_empty = False
    for line in lines:
        if line.strip():
            cleaned_lines.append(line)
            prev_empty = False
        elif not prev_empty:
            cleaned_lines.append('')
            prev_empty = True
    
    text_content = '\n'.join(cleaned_lines).strip()
    
    # If still empty, try alternative extraction methods
    if not text_content.strip():
        # Method 1: Get all text using python-docx's built-in method
        try:
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text)
            text_content = '\n'.join(all_text).strip()
        except Exception:
            pass
        
        # Method 2: Try reading raw XML
        if not text_content.strip():
            try:
                import xml.etree.ElementTree as ET
                # Get the document XML
                xml_str = doc._body._body.xml
                root = ET.fromstring(xml_str)
                # Extract all text nodes
                text_nodes = []
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        text_nodes.append(elem.text.strip())
                if text_nodes:
                    text_content = '\n'.join(text_nodes)
            except Exception:
                pass
    
    # Final check
    if not text_content.strip():
        raise ValueError(
            f"DOCX file appears to be empty or contains no extractable text. "
            f"Found {para_count} paragraphs and {table_count} tables. "
            f"The document may contain only images, shapes, or unsupported content."
        )
    
    return text_content
