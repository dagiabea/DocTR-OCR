# Copyright (C) 2021-2025, Mindee.
# This program is licensed under the Apache License 2.0.
# See LICENSE or go to <https://opensource.org/licenses/Apache-2.0> for full license details.

import io
from typing import Any

import numpy as np
from PIL import Image, ImageDraw, ImageFont

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    import striprtf
    RTF_AVAILABLE = True
except ImportError:
    striprtf = None
    RTF_AVAILABLE = False

try:
    from odf import text, teletype
    from odf.opendocument import load
    ODT_AVAILABLE = True
except ImportError:
    load = None
    teletype = None
    ODT_AVAILABLE = False


def text_to_image(text: str, width: int = 1200, padding: int = 50) -> np.ndarray:
    """Convert text to an image for OCR processing.
    
    Args:
        text: Text content to convert
        width: Image width in pixels
        padding: Padding around text in pixels
        
    Returns:
        numpy array representing the image
    """
    # Create image with white background
    lines = text.split('\n')
    
    # Estimate height (rough calculation)
    line_height = 30
    height = len(lines) * line_height + (padding * 2)
    
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except (OSError, IOError):
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        except (OSError, IOError):
            font = ImageFont.load_default()
    
    y = padding
    for line in lines:
        if line.strip():
            draw.text((padding, y), line, fill='black', font=font)
        y += line_height
    
    # Convert PIL Image to numpy array
    return np.array(img)


def convert_docx_to_images(file_content: bytes) -> list[np.ndarray]:
    """Convert DOCX file to list of images (one per page/section).
    
    Args:
        file_content: Binary content of DOCX file
        
    Returns:
        List of numpy arrays representing pages
    """
    if not DOCX_AVAILABLE or DocxDocument is None:
        import sys
        python_path = sys.executable
        raise ValueError(
            f"python-docx is required for DOCX support. "
            f"Install it with: {python_path} -m pip install python-docx\n"
            f"Current Python: {python_path}"
        )
    
    doc = DocxDocument(io.BytesIO(file_content))
    
    # Extract all text from paragraphs
    full_text = []
    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if para_text:  # Only add non-empty paragraphs
            full_text.append(para_text)
    
    # Extract text from tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_rows.append(' | '.join(row_text))
        if table_rows:
            full_text.append('')  # Add blank line before table
            full_text.extend(table_rows)
            full_text.append('')  # Add blank line after table
    
    # Also try to extract from headers and footers
    try:
        for section in doc.sections:
            if section.header:
                header_text = []
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text.strip())
                if header_text:
                    full_text.insert(0, '\n'.join(header_text))
                    full_text.insert(1, '')  # Blank line after header
            
            if section.footer:
                footer_text = []
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text.strip())
                if footer_text:
                    full_text.append('')  # Blank line before footer
                    full_text.extend(footer_text)
    except Exception:
        # Headers/footers extraction is optional, continue if it fails
        pass
    
    text_content = '\n'.join(full_text)
    
    # If still empty, try getting all text using a different method
    if not text_content.strip():
        # Fallback: get all text content
        try:
            text_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception:
            pass
    
    # Convert to image
    if text_content.strip():
        return [text_to_image(text_content)]
    else:
        raise ValueError("DOCX file appears to be empty or contains no extractable text")


def convert_doc_to_images(file_content: bytes) -> list[np.ndarray]:
    """Convert DOC file to images.
    
    Note: Old .doc format is harder to parse. This is a basic implementation.
    For better support, consider using antiword or similar tools.
    
    Args:
        file_content: Binary content of DOC file
        
    Returns:
        List of numpy arrays representing pages
    """
    # DOC format is binary and complex. For now, raise an error suggesting conversion
    raise ValueError(
        "Direct .doc file support is limited. Please convert to DOCX or PDF first. "
        "You can use online converters or LibreOffice to convert DOC to DOCX."
    )


def convert_txt_to_images(file_content: bytes, encoding: str = 'utf-8') -> list[np.ndarray]:
    """Convert TXT file to images.
    
    Args:
        file_content: Binary content of TXT file
        encoding: Text encoding (default: utf-8)
        
    Returns:
        List of numpy arrays representing pages
    """
    try:
        text_content = file_content.decode(encoding)
    except UnicodeDecodeError:
        # Try other common encodings
        for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
            try:
                text_content = file_content.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(f"Could not decode text file. Tried encodings: utf-8, latin-1, cp1252, iso-8859-1")
    
    if text_content.strip():
        return [text_to_image(text_content)]
    else:
        raise ValueError("TXT file appears to be empty")


def convert_rtf_to_images(file_content: bytes) -> list[np.ndarray]:
    """Convert RTF file to images.
    
    Args:
        file_content: Binary content of RTF file
        
    Returns:
        List of numpy arrays representing pages
    """
    if striprtf is None:
        raise ValueError("striprtf is required for RTF support. Install it with: pip install striprtf")
    
    try:
        rtf_text = file_content.decode('utf-8', errors='ignore')
        plain_text = striprtf.striprtf.rtf_to_text(rtf_text)
    except Exception as e:
        raise ValueError(f"Failed to parse RTF file: {str(e)}")
    
    if plain_text.strip():
        return [text_to_image(plain_text)]
    else:
        raise ValueError("RTF file appears to be empty")


def convert_odt_to_images(file_content: bytes) -> list[np.ndarray]:
    """Convert ODT file to images.
    
    Args:
        file_content: Binary content of ODT file
        
    Returns:
        List of numpy arrays representing pages
    """
    if load is None or teletype is None:
        raise ValueError("odfpy is required for ODT support. Install it with: pip install odfpy")
    
    try:
        doc = load(io.BytesIO(file_content))
        paragraphs = doc.getElementsByType(text.P)
        
        full_text = []
        for paragraph in paragraphs:
            full_text.append(teletype.extractText(paragraph))
        
        text_content = '\n'.join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse ODT file: {str(e)}")
    
    if text_content.strip():
        return [text_to_image(text_content)]
    else:
        raise ValueError("ODT file appears to be empty")


def convert_document_to_images(file_content: bytes, mime_type: str, filename: str | None = None) -> list[np.ndarray]:
    """Convert various document formats to images for OCR processing.
    
    Args:
        file_content: Binary content of the file
        mime_type: MIME type of the file
        filename: Optional filename for better format detection
        
    Returns:
        List of numpy arrays representing pages
        
    Raises:
        ValueError: If format is unsupported or conversion fails
    """
    # Map MIME types to conversion functions
    mime_to_converter = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': convert_docx_to_images,  # .docx
        'application/msword': convert_doc_to_images,  # .doc
        'text/plain': convert_txt_to_images,  # .txt
        'application/rtf': convert_rtf_to_images,  # .rtf
        'application/vnd.oasis.opendocument.text': convert_odt_to_images,  # .odt
    }
    
    # Also check filename extension as fallback
    if filename:
        filename_lower = filename.lower()
        if filename_lower.endswith('.docx'):
            return convert_docx_to_images(file_content)
        elif filename_lower.endswith('.doc'):
            return convert_doc_to_images(file_content)
        elif filename_lower.endswith('.txt'):
            return convert_txt_to_images(file_content)
        elif filename_lower.endswith('.rtf'):
            return convert_rtf_to_images(file_content)
        elif filename_lower.endswith('.odt'):
            return convert_odt_to_images(file_content)
    
    # Use MIME type
    converter = mime_to_converter.get(mime_type)
    if converter:
        return converter(file_content)
    
    raise ValueError(f"Unsupported file format: {mime_type} for file {filename}")
